"""
CV parsing service.

Two steps:
1. Extract raw text from PDF using PyMuPDF (fast, no API call)
2. Send raw text to LLM → returns structured CV as JSON

The structured JSON is what gets stored in MongoDB and used
for rating + CV tailoring in Week 2.
"""

import json
import re

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from odf import teletype
from odf.opendocument import load as load_odf
from langchain_core.messages import HumanMessage, SystemMessage
from langsmith import traceable

from config import settings
from services.ai_models import get_cost_multiplier, get_default_model_for_provider
from services.ai_usage import record_from_llm_response
from services.llm import get_llm
from services.prompt_safety import fence


# ── Step 1: file bytes → raw text (format-specific) ──────
def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages = []
    link_notes = []
    for page in doc:
        pages.append(page.get_text())
        # get_text() only returns visible text, a hyperlinked label like "LinkedIn"
        # carries its URL as a link annotation, not in the text stream, so the LLM
        # parser never sees it unless it's spliced in separately here.
        for link in page.get_links():
            uri = link.get("uri")
            if link.get("kind") == fitz.LINK_URI and uri:
                label = page.get_textbox(link["from"]).strip()
                if label:
                    link_notes.append(f"[{label}: {uri}]")
    doc.close()
    raw = "\n".join(pages).strip()
    if not raw:
        raise ValueError("PDF appears to be empty or scanned (no extractable text).")
    if link_notes:
        raw += "\n\n" + "\n".join(link_notes)
    return raw


def extract_text_from_docx(file_bytes: bytes) -> str:
    import io

    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.extend(cell.text for cell in row.cells if cell.text.strip())
    raw = "\n".join(paragraphs).strip()
    if not raw:
        raise ValueError("Word document appears to be empty.")
    return raw


def extract_text_from_odt(file_bytes: bytes) -> str:
    import io

    doc = load_odf(io.BytesIO(file_bytes))
    raw = teletype.extractText(doc.text).strip()
    if not raw:
        raise ValueError("ODT document appears to be empty.")
    return raw


def extract_text_from_plain(file_bytes: bytes) -> str:
    """Covers .txt and .tex, both are just plain text to the parser."""
    raw = file_bytes.decode("utf-8", errors="ignore").strip()
    if not raw:
        raise ValueError("File appears to be empty.")
    return raw


_EXTRACTORS = {
    "pdf": extract_text_from_pdf,
    "docx": extract_text_from_docx,
    "odt": extract_text_from_odt,
    "txt": extract_text_from_plain,
    "tex": extract_text_from_plain,
}


def extract_text(file_bytes: bytes, extension: str) -> str:
    extractor = _EXTRACTORS.get(extension.lower().lstrip("."))
    if not extractor:
        raise ValueError(f"Unsupported CV file format: .{extension}")
    return extractor(file_bytes)


# ── Step 2: raw text → structured JSON via LLM ──────────
SYSTEM_PROMPT = """
You are a CV parser. Extract structured information from the CV text provided.

Return ONLY valid JSON, no markdown, no backticks, no explanation.

The JSON must follow this exact structure:
{
  "name": "string",
  "email": "string or null",
  "phone": "string or null",
  "location": "string or null",
  "summary": "string, the professional summary or objective",
  "links": {
    "website": "personal site/portfolio URL exactly as written, or null",
    "github": "GitHub profile URL exactly as written, or null",
    "linkedin": "LinkedIn profile URL exactly as written, or null"
  },
  "skills": [
    {"category": "string, e.g. Backend, Frontend, Cloud & DevOps, Machine Learning", "items": ["skill1", "skill2"]}
  ],
  "experience": [
    {
      "title": "string",
      "company": "string, the organization name only, never append an event/context description with a pipe or dash",
      "start": "string e.g. 2022, or null if this was a one-off event with no real duration",
      "end": "string e.g. 2025 or Present, or null if this was a one-off event with no real duration",
      "bullets": ["bullet1", "bullet2"]
    }
  ],
  "projects": [
    {
      "name": "string",
      "description": "string, one sentence summary",
      "tech": ["tech1", "tech2"],
      "bullets": ["bullet1", "bullet2"],
      "live_url": "deployed/demo link exactly as written, or null",
      "repo_url": "source code/GitHub repo link exactly as written, or null"
    }
  ],
  "education": [
    {
      "degree": "string",
      "institution": "string",
      "start": "string",
      "end": "string",
      "grade": "string or null",
      "specialization": "string or null, a Minor/Honours/concentration explicitly listed for this degree, e.g. 'Minor (Honours) in AI & ML in Healthcare', null if none is written"
    }
  ],
  "languages": ["English", "etc"],
  "certifications": []
}

Rules:
- Extract ONLY what is actually in the CV. Never invent or assume.
- If a field is missing, use null for strings or [] for arrays.
- skills should be individual technologies/tools, not sentences, grouped under the
  candidate's own natural categories (or close to them). Put anything that doesn't fit
  a clear category under "Other" rather than forcing it somewhere it doesn't belong.
- Keep bullet points concise and exactly as written in the CV.
- links/live_url/repo_url: only fill these in if a URL is actually written in the CV
  text, verbatim, never construct or guess one (e.g. never invent a github.com/username
  URL just because a username is mentioned elsewhere). null if not present. Hyperlinked
  labels appear as "[Label: URL]" notes near the end of the text, e.g. "[LinkedIn:
  https://...]", match those to the right field by label the same as any other URL
  written in the CV.
- specialization: capture a Minor, Honours designation, or concentration ONLY if the CV
  explicitly names one for that specific degree, as its own field, not folded into
  "degree". null if the degree line doesn't mention one, never guess or infer.
- The contact block has been replaced with [REDACTED_PHONE] / [REDACTED_EMAIL]
  placeholders, leave "phone" and "email" as null, they're filled in locally.
""".strip()


def flatten_skills(skills: list) -> list[str]:
    """skills is either the categorized shape ([{"category": ..., "items": [...]}]) or
    the older flat list[str] shape (CVs parsed before categorization existed), flattens
    either into a plain list of skill names for callers that just need the names, not
    the grouping (rating prompts, keyword matching, skill counts)."""
    if not skills:
        return []
    if isinstance(skills, dict):
        skills = [skills]
    if isinstance(skills[0], dict):
        return [
            s
            for group in skills
            if isinstance(group, dict)
            for s in (group.get("items") or [])
            if isinstance(s, str)
        ]
    return [s for s in skills if isinstance(s, str)]


# Search queries, not the full CV skill dump. Long parenthetical names
# ("GANs (Generative Adversarial Networks, ...)") make Indeed/Jooble return 0.
_SEARCH_SKILL_CAP = 8
_SEARCH_SKILL_MAX_LEN = 32
_SEARCH_SKILL_SKIP = {
    "html5",
    "css3",
    "git",
    "github",
    "git / github",
    "postman",
    "agile",
    "agile / scrum",
    "scrum",
    "performance optimisation",
    "technical documentation",
}


def search_skills_from_cv(skills: list) -> list[str]:
    """Short skill names for job-board search queries. One lead skill per
    category when grouped, else first N short names from a flat list."""
    if not skills:
        return []
    if isinstance(skills, dict):
        skills = [skills]
    picked: list[str] = []
    seen: set[str] = set()

    def consider(name: str) -> None:
        s = name.strip()
        if not s or len(s) > _SEARCH_SKILL_MAX_LEN:
            return
        key = s.lower()
        if key in seen or key in _SEARCH_SKILL_SKIP:
            return
        seen.add(key)
        picked.append(s)

    if isinstance(skills[0], dict):
        for group in skills:
            if not isinstance(group, dict) or len(picked) >= _SEARCH_SKILL_CAP:
                break
            for item in group.get("items") or []:
                if isinstance(item, str):
                    consider(item)
                    if item.strip() in picked:
                        break
    else:
        for item in skills:
            if len(picked) >= _SEARCH_SKILL_CAP:
                break
            if isinstance(item, str):
                consider(item)
    return picked[:_SEARCH_SKILL_CAP]


_PHONE_RE = re.compile(r"(\+?\d[\d\-.\s()]{7,}\d)")
_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_MINOR_HONOURS_RE = re.compile(
    r"\b(minor|honours|honors|concentration)\b", re.IGNORECASE
)


def _redact_contact_details(text: str) -> str:
    """Mask phone numbers and emails before the text leaves the server."""
    text = _EMAIL_RE.sub("[REDACTED_EMAIL]", text)
    text = _PHONE_RE.sub("[REDACTED_PHONE]", text)
    return text


def _extract_contact_details(text: str) -> tuple[str | None, str | None]:
    """Pull the real phone/email locally, never sent to the LLM."""
    email_match = _EMAIL_RE.search(text)
    phone_match = _PHONE_RE.search(text)
    return (
        phone_match.group(0).strip() if phone_match else None,
        email_match.group(0) if email_match else None,
    )


@traceable(name="parse_cv_with_llm", run_type="llm")
async def parse_cv_with_llm(raw_text: str, user: dict | None = None) -> dict:
    user = user or {}
    user_id = str(user.get("_id", "")) or None
    user_provider = user.get("cv_parsing_provider") or None
    user_model = user.get("cv_parsing_model") or (
        await get_default_model_for_provider(user_provider, "cv_parsing")
        if user_provider
        else None
    )
    cost_multiplier = await get_cost_multiplier(user_provider, user_model, "cv_parsing")
    llm = get_llm(provider=user_provider, model=user_model)
    provider = user_provider or settings.llm_provider

    redacted_text = _redact_contact_details(raw_text)
    messages = [
        SystemMessage(content=SYSTEM_PROMPT),
        HumanMessage(content=f"Parse this CV:\n\n{fence('CV TEXT', redacted_text)}"),
    ]

    try:
        response = await llm.ainvoke(messages)
    except Exception as e:
        raise ValueError(f"CV parsing failed ({provider}): {e}") from e
    model = getattr(
        llm, "model", getattr(llm, "model_name", user_model or settings.openai_model)
    )
    model = str(model or "unknown")
    if user_id:
        await record_from_llm_response(
            user_id,
            response,
            operation="cv_parse",
            provider=provider,
            model=model,
            cost_multiplier=cost_multiplier,
        )
    content = response.content.strip()

    # strip markdown fences if model wraps anyway
    content = re.sub(r"^```[a-z]*\n?", "", content)
    content = re.sub(r"\n?```$", "", content)
    content = content.strip()

    try:
        parsed = json.loads(content)
    except json.JSONDecodeError as e:
        raise ValueError(f"LLM returned invalid JSON: {e}\n\nRaw output:\n{content}")

    phone, email = _extract_contact_details(raw_text)
    parsed["phone"] = phone
    parsed["email"] = email
    parsed["parsed_by_model"] = f"{provider}:{model}"

    # Deterministic backstop, not another prompt ask: the LLM can miss a minor/honours
    # line, this can't recover the text, but it can at least flag the CV for review
    # instead of the gap staying invisible (surfaced via Settings' existing CV review
    # banner for auto-filled fields).
    education = parsed.get("education") or []
    has_specialization = any((edu or {}).get("specialization") for edu in education)
    if education and not has_specialization and _MINOR_HONOURS_RE.search(raw_text):
        parsed["parse_warnings"] = [
            "Your CV mentions a Minor/Honours/concentration but we couldn't confidently "
            "extract it into the Education section, please check and add it manually if needed."
        ]

    return parsed


# ── Combined: file bytes → structured dict ───────────────
async def process_cv(
    file_bytes: bytes, extension: str = "pdf", user: dict | None = None
) -> tuple[str, dict]:
    """
    Returns (raw_text, structured_json).
    We store both, raw_text for embedding later, structured for display.
    """
    raw_text = extract_text(file_bytes, extension)
    structured = await parse_cv_with_llm(raw_text, user=user)
    return raw_text, structured
